import os
import docx
import PyPDF2
from docx.oxml.ns import qn

def allowed_file(filename, allowed_extensions):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in allowed_extensions

def extract_text_from_file(filepath):
    """
    【带坐标版】文本提取工具
    1. PDF：自动插入 【第X页】 标记。
    2. Word：保留大纲结构，辅助大模型定位章节。
    """
    _, ext_with_dot = os.path.splitext(filepath)
    ext = ext_with_dot.lstrip('.').lower()
    
    text = ""
    
    try:
        if ext == 'docx':
            doc = docx.Document(filepath)
            
            # Word 文档没有固定页码，我们尽可能提取“章节标题”作为定位符
            # 同时为了不破坏上下文，我们保留段落结构
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    # 如果是标题样式，或者看起来像标题（短且不含标点），可以特殊标记
                    if para.style.name.startswith('Heading'):
                        text += f"\n【章节：{para.text}】\n"
                    else:
                        text += para.text + "\n"
            
            # 表格内容
            for table in doc.tables:
                text += "\n【表格数据区域】\n"
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        text += " | ".join(row_data) + "\n"
            
            # 文本框/XML 深度扫描 (保留之前的功能，但标记为非正文)
            if len(text) < 200: 
                print("【系统提示】检测到正文内容较少，正在启动深度XML扫描...")
                xml_text = []
                for element in doc.element.body.iter():
                    if element.tag.endswith('}t'): 
                        if element.text:
                            xml_text.append(element.text)
                if len("".join(xml_text)) > len(text):
                    text = "【深度扫描内容（位置未知）】\n" + "\n".join(xml_text)

        elif ext == 'pdf':
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                count = 0
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        # 【核心修改】插入页码标记，让大模型知道这是第几页
                        text += f"\n\n--- 【第 {i+1} 页】 ---\n\n"
                        text += page_text
                        count += 1
                
                if len(text.strip()) < 50 and count > 0:
                    return "【错误】检测到这可能是‘扫描版PDF’。本系统暂不支持OCR，请上传 Word 文档。"

        elif ext == 'txt':
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
        
        else:
            return f"不支持的文件类型: {ext}"

    except Exception as e:
        print(f"解析异常: {str(e)}")
        return f"文件读取失败: {str(e)}"
        
    if len(text.strip()) < 10:
        return f"【警告】文件内容为空或无法识别。"

    return text[:15000]